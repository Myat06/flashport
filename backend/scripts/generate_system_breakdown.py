"""
Generate FlashPort System Breakdown as a Word (.docx) document.
Output: docs/FlashPort_System_Breakdown_Updated.docx

Usage:
    cd backend && source venv/bin/activate
    python scripts/generate_system_breakdown.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent.parent.parent / "docs" / "FlashPort_System_Breakdown_Updated.docx"

# ── Colours ───────────────────────────────────────────────────────────────────
BLUE    = RGBColor(0x1B, 0x4F, 0xBF)
DBLUE   = RGBColor(0x0F, 0x2E, 0x7A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY    = RGBColor(0xF5, 0xF5, 0xF5)
DGRAY   = RGBColor(0x37, 0x41, 0x51)
GREEN   = RGBColor(0x16, 0xA3, 0x4A)
RED     = RGBColor(0xDC, 0x26, 0x26)
YELLOW  = RGBColor(0xCA, 0x8A, 0x04)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"] if level == 1 else doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = DBLUE if level == 1 else BLUE
    run.font.size = Pt(16 if level == 1 else 13)
    run.bold = True
    return p


def add_body(doc, text, bold=False, color=None, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell_text(hdr.cells[i], h, bold=True, color=WHITE, size=8.5)
        set_cell_bg(hdr.cells[i], "0F2E7A")

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell_text(tr.cells[ci], val, size=8.5)
            set_cell_bg(tr.cells[ci], bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = DGRAY


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FlashPort")
    r.bold = True; r.font.size = Pt(32); r.font.color.rgb = DBLUE

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("System Breakdown — Updated June 2026")
    r2.font.size = Pt(14); r2.font.color.rgb = BLUE

    doc.add_paragraph()
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("AI Open Innovation Challenge 2026  |  Case 1: Cikarang Dry Port")
    r3.font.size = Pt(11); r3.font.color.rgb = DGRAY

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("Team: Teknik Logistik  |  President University  |  Final Deadline: 31 August 2026")
    r4.font.size = Pt(10); r4.font.color.rgb = DGRAY

    doc.add_page_break()

    # ── SECTION 1: OVERVIEW ──────────────────────────────────────────────────
    add_heading(doc, "1. System Overview")
    add_body(doc,
        "FlashPort automates customs declaration preparation for Cikarang Dry Port (CDP) "
        "by combining mobile document capture, a multi-stage AI pipeline, and a professional "
        "manager dashboard. The system eliminates manual data entry, pre-scores declarations "
        "for CEISA rejection risk, and gives managers full visibility before submission.")
    doc.add_paragraph()

    add_heading(doc, "1.1 Problem vs Solution", level=2)
    add_table(doc,
        ["Stage", "Before FlashPort", "After FlashPort"],
        [
            ["Data collection", "Manual transcription from paper", "Mobile camera / file upload"],
            ["Field extraction", "Human reading + typing", "spaCy NER deep learning + Regex"],
            ["Risk assessment", "None", "XGBoost AI scorer + SHAP explainability"],
            ["Manager review", "Manual email chain", "Professional web dashboard"],
            ["CEISA submission", "Manual portal entry", "One-click or batch from dashboard"],
            ["Processing time", "2-4 hours per declaration", "Under 3 minutes"],
        ],
        col_widths=[1.5, 2.3, 2.3]
    )

    # ── SECTION 2: ARCHITECTURE ──────────────────────────────────────────────
    add_heading(doc, "2. System Architecture")
    add_body(doc, "FlashPort has three independent layers that communicate via REST API and WebSocket.")
    doc.add_paragraph()

    add_heading(doc, "2.1 Three-Layer Architecture", level=2)
    add_table(doc,
        ["Layer", "Technology", "Key Responsibilities"],
        [
            ["Field Layer\n(Mobile App)", "Flutter 3.x (iOS + Android)",
             "Offline-first document capture, SQLite local storage, FCM push notifications, status timeline"],
            ["Processing Layer\n(Backend)", "FastAPI + Python 3.11",
             "PDF/image ingestion, OCR, NER field extraction, XGBoost risk scoring, SHAP explainability, PostgreSQL storage"],
            ["Management Layer\n(Web Dashboard)", "React 18 + Tailwind CSS",
             "9-page pro dashboard: declarations, approve/reject, CEISA submit, operators, watchlist, risk rules, SLA, audit"],
        ],
        col_widths=[1.5, 1.8, 2.8]
    )

    doc.add_paragraph()
    add_heading(doc, "2.2 Technology Stack", level=2)
    add_table(doc,
        ["Component", "Technology", "Notes"],
        [
            ["Mobile", "Flutter 3.x (Dart)", "iOS + Android from single codebase"],
            ["Mobile storage", "SQLite (sqflite)", "Offline-first local persistence"],
            ["Backend", "Python 3.11 + FastAPI", "Async, 36 API endpoints, OpenAPI docs"],
            ["PDF text extraction", "pypdf", "Direct text from digital PDFs — no OCR needed"],
            ["OCR engine", "Tesseract eng+ind", "For scanned/photo documents only"],
            ["Image preprocessing", "OpenCV (smart pipeline)", "Skips aggressive processing for clean images"],
            ["Field extraction — Stage 1", "spaCy NER (custom trained)", "11 entity types, 100% F1, English + Indonesian"],
            ["Field extraction — Stage 2", "Python Regex fallback", "Fills NER gaps, table-merge OCR patterns"],
            ["Risk scoring", "XGBoost (300 trees)", "89.5% cross-validation accuracy"],
            ["Explainability", "SHAP TreeExplainer", "16 per-feature contributions per prediction"],
            ["Database", "PostgreSQL 15", "Local install, no Docker required"],
            ["Web frontend", "React 18 + Tailwind CSS 3", "Sidebar layout, 9 pages"],
            ["Push notifications", "Firebase FCM HTTP v1", "Approve/reject + CEISA lane results"],
            ["Auth", "JWT HS256 + API Key", "8h manager sessions, operator PIN login"],
        ],
        col_widths=[1.8, 1.8, 2.5]
    )

    doc.add_page_break()

    # ── SECTION 3: MOBILE ────────────────────────────────────────────────────
    add_heading(doc, "3. Mobile Application")

    add_heading(doc, "3.1 Offline-First Upload Flow", level=2)
    steps = [
        "Operator logs in with Employee ID + PIN",
        "Select document type: Commercial Invoice / Bill of Lading / Packing List",
        "Take photo (camera) OR attach file (JPG, PNG, PDF from Files app)",
        "Preview screen shows: document image with green Attached badge, file name + size",
        "Tap Upload & Sync — 3-step progress animation plays",
        "Online: Result screen shows risk score, extracted fields, CEISA readiness",
        "Offline: Saved as Pending locally, auto-syncs when network detected",
        "After sync: status timeline updates Scanned -> Synced -> Reviewed -> Approved/Rejected",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"Step {i}: {s}")

    doc.add_paragraph()
    add_heading(doc, "3.2 Status Timeline on Scan Tile", level=2)
    add_body(doc, "Each scan in the history list shows a mini progress timeline:")
    add_code(doc, ["Scanned  ->  Synced  ->  Reviewed  ->  Approved",
                   "                                    ->  Rejected"])

    doc.add_paragraph()
    add_heading(doc, "3.3 Push Notifications (FCM)", level=2)
    add_table(doc,
        ["Event", "Notification Title", "Body"],
        [
            ["Manager approves", "Declaration Approved", "Manager note (if provided)"],
            ["Manager rejects", "Declaration Rejected", "Manager note with reason"],
            ["CEISA Green Lane", "CEISA: Green Lane", "Reference number — immediate release"],
            ["CEISA Yellow Lane", "CEISA: Yellow Lane", "Reference number — document check"],
            ["CEISA Red Lane", "CEISA: Red Lane", "Reference number — physical inspection"],
        ],
        col_widths=[1.5, 1.8, 2.8]
    )

    doc.add_paragraph()
    add_heading(doc, "3.4 Physical Device Setup", level=2)
    add_body(doc, "For physical iOS/Android devices the Server URL must be set to the Mac's local IP address.")
    add_body(doc, "Find IP: ipconfig getifaddr en0", bold=True)
    add_body(doc, "Set in app login screen Server URL field: http://[YOUR-IP]:8000")
    add_body(doc, "iOS Note: NSAllowsArbitraryLoads = true is already set in Info.plist to allow local HTTP.")

    doc.add_page_break()

    # ── SECTION 4: BACKEND PIPELINE ──────────────────────────────────────────
    add_heading(doc, "4. Backend Processing Pipeline")

    add_heading(doc, "4.1 Two-Path Document Ingestion", level=2)
    add_table(doc,
        ["Input type", "Detection", "Processing path"],
        [
            ["Digital PDF", "base64 starts with JVBERi0 (= %PDF-)",
             "pypdf extracts text directly from text layer. No OCR needed. Best accuracy for generated/digital PDFs."],
            ["Scanned / image PDF", "Is PDF but pypdf returns < 50 chars",
             "Fall back to pdf2image conversion then Tesseract OCR on each page."],
            ["Phone photo (JPG/PNG)", "base64 starts with /9j/ or iVBORw",
             "Smart OpenCV preprocessing then Tesseract OCR (eng+ind)."],
        ],
        col_widths=[1.4, 1.8, 2.9]
    )

    doc.add_paragraph()
    add_heading(doc, "4.2 Smart Image Preprocessing", level=2)
    add_body(doc,
        "Not all images need aggressive preprocessing. The system detects image quality "
        "using the white pixel ratio and applies the appropriate pipeline:")
    add_table(doc,
        ["Image quality", "White pixel ratio", "Processing applied"],
        [
            ["Clean digital image or well-lit photo", "> 55% white pixels",
             "Grayscale only — Tesseract reads it perfectly as-is"],
            ["Dark, shadowed, or noisy photo", "< 55% white pixels",
             "Grayscale + Gaussian denoise + adaptive threshold + deskew rotation"],
        ],
        col_widths=[2.2, 1.6, 2.3]
    )

    doc.add_paragraph()
    add_heading(doc, "4.3 Two-Stage NLP Field Extraction", level=2)
    add_body(doc, "Stage 1 — spaCy NER (primary extractor):", bold=True)
    add_bullet(doc, "Custom trained model on 300 annotated customs document examples")
    add_bullet(doc, "11 entity types: HS_CODE, INVOICE_VALUE, CONTAINER_ID, IMPORTER, EXPORTER, NET_WEIGHT, GROSS_WEIGHT, VESSEL_NAME, PORT_OF_ORIGIN, INVOICE_NUMBER, CARTON_COUNT")
    add_bullet(doc, "Handles both English labels (CONSIGNEE, SHIPPER) and Indonesian (PENERIMA, PENGIRIM)")
    add_bullet(doc, "100% F1 on development set. Early stopped at iteration 15.")
    doc.add_paragraph()
    add_body(doc, "Stage 2 — Regex fallback:", bold=True)
    add_bullet(doc, "Runs only on fields NER missed")
    add_bullet(doc, "Specialized patterns for OCR table-merge: catches 'Apple Inc. PT Importer' on same line")
    add_bullet(doc, "TOTAL row pattern for packing list weights, direct MV vessel pattern")
    add_bullet(doc, "Validation layer: rejects dates as weights, bare currency codes, invalid container formats")

    doc.add_paragraph()
    add_heading(doc, "4.4 Extracted Fields", level=2)
    add_table(doc,
        ["Field", "Extraction pattern", "Example"],
        [
            ["HS Code", "8-digit XX.XX.XX pattern", "8471.30.00"],
            ["Invoice Value", "Currency code + number (TOTAL row preferred)", "USD 133,800.94"],
            ["Container ID", "Exactly 4 letters + 7 digits", "YMLU7593824"],
            ["Importer", "CONSIGNEE/PENERIMA context + PT prefix", "PT LG Electronics Indonesia"],
            ["Exporter", "SHIPPER/PENGIRIM context", "Apple Inc."],
            ["Net Weight", "NET WEIGHT label + KG, or TOTAL row", "460.59 KG"],
            ["Gross Weight", "GROSS WEIGHT label + KG, or TOTAL row", "497.44 KG"],
            ["Vessel Name", "MV prefix pattern", "MV Yang Ming Wellness"],
            ["Port of Origin", "PORT OF LOADING context", "Rotterdam, Netherlands"],
            ["Invoice Number", "INV/SI/PO prefix", "INV-2026-5969"],
            ["Carton Count", "Number + CARTONS/CTN", "7 CTN"],
        ],
        col_widths=[1.5, 2.5, 2.1]
    )

    doc.add_page_break()

    # ── SECTION 5: RISK SCORING ──────────────────────────────────────────────
    add_heading(doc, "5. AI Risk Scoring — XGBoost + SHAP")

    add_heading(doc, "5.1 XGBoost Model", level=2)
    add_body(doc, "The model uses 16 features and outputs probabilities for Green/Yellow/Red lane:")
    add_table(doc,
        ["Feature", "Importance", "Why it matters"],
        [
            ["is_restricted_hs", "36.5%", "Weapons/chemicals HS code — strongest signal by far"],
            ["hs_high_scrutiny", "10.7%", "Electronics, petroleum — moderate scrutiny category"],
            ["missing_field_count", "8.0%", "Total critical fields missing"],
            ["has_port", "6.0%", "Port of origin present or not"],
            ["has_container_id", "4.6%", "Container ID present or not"],
            ["high_value_no_container", "4.6%", "High value (>USD 50k) without container ID"],
            ["has_hs_code", "4.5%", "HS code present or not"],
            ["has_vessel", "4.0%", "Vessel name present or not"],
            ["invoice_value_log", "4.0%", "Log-scaled invoice value"],
            ["confidence_score", "2.8%", "OCR quality: high=2, medium=1, low=0"],
            ["is_high_value", "2.2%", "Invoice > USD 50,000"],
            ["is_very_high_value", "2.7%", "Invoice > USD 200,000"],
            ["... 4 more features", "8.5%", "has_importer, has_exporter, document_type, has_invoice_value"],
        ],
        col_widths=[1.8, 1.0, 3.3]
    )

    doc.add_paragraph()
    add_heading(doc, "5.2 Risk Score to Lane Mapping", level=2)
    add_table(doc,
        ["Score range", "Lane", "Action", "CEISA code"],
        [
            ["0 - 29%", "Green Lane", "Immediate release — document accepted", "SP2-200"],
            ["30 - 69%", "Yellow Lane", "Document verification required", "SP2-412"],
            ["70 - 100%", "Red Lane", "Physical inspection required", "SP2-500"],
        ],
        col_widths=[1.0, 1.2, 2.8, 1.1]
    )

    doc.add_paragraph()
    add_heading(doc, "5.3 Additional Risk Boosts", level=2)
    add_bullet(doc, "Watchlist hit (importer or exporter flagged by manager): +25 points per hit")
    add_bullet(doc, "Manager-configured DB rules (custom field/condition/boost): applied on every scan if active")
    add_bullet(doc, "Watchlist and rules are stacked on top of the XGBoost base score")

    doc.add_paragraph()
    add_heading(doc, "5.4 SHAP Explainability", level=2)
    add_body(doc,
        "Every prediction includes 16 SHAP values showing exactly which features pushed the "
        "risk score up or down. These are displayed as a bar chart in the web dashboard detail "
        "panel so the manager can understand why a declaration was flagged.")

    doc.add_page_break()

    # ── SECTION 6: AI TRAINING ───────────────────────────────────────────────
    add_heading(doc, "6. AI Training Pipeline")

    add_heading(doc, "6.1 XGBoost Risk Model Training", level=2)
    add_body(doc, "Step 1 — Generate synthetic training data:", bold=True)
    add_code(doc, ["python scripts/generate_training_data.py",
                   "# Output: data/training_declarations.csv",
                   "# 600 records: 395 Green / 127 Yellow / 78 Red"])
    doc.add_paragraph()
    add_table(doc,
        ["Profile", "Count", "Description", "Label"],
        [
            ["Clean", "330 (55%)", "All fields present, valid HS, high confidence", "Green"],
            ["Partial", "150 (25%)", "1-2 non-critical fields missing", "Yellow"],
            ["Missing", "72 (12%)", "2-4 critical fields missing (HS, importer, value, container)", "Red"],
            ["Risky", "48 (8%)", "Restricted HS code, very high value, low OCR confidence", "Red"],
        ],
        col_widths=[1.2, 1.0, 3.0, 0.9]
    )
    doc.add_paragraph()
    add_body(doc, "Step 2 — Train the model:", bold=True)
    add_code(doc, ["python scripts/train_risk_model.py",
                   "# 5-fold cross-validation: 89.5% accuracy",
                   "# Output: models/risk_model.xgb + models/model_info.json"])

    doc.add_paragraph()
    add_heading(doc, "6.2 spaCy NER Model Training", level=2)
    add_body(doc, "Step 1 — Generate annotated examples:", bold=True)
    add_code(doc, ["python scripts/generate_ner_training.py",
                   "# Output: data/ner_training.spacy",
                   "# 300 examples, English + Indonesian labels"])
    doc.add_paragraph()
    add_body(doc, "Step 2 — Train the NER model:", bold=True)
    add_code(doc, ["python scripts/train_ner_model.py",
                   "# Base: en_core_web_sm, fine-tuned on customs documents",
                   "# Early stop at iter 15 — F1: 1.000 on all 11 entities",
                   "# Output: models/ner_model/"])

    doc.add_paragraph()
    add_heading(doc, "6.3 Sample Document Generation", level=2)
    add_code(doc, ["python scripts/generate_sample_docs.py",
                   "# Output: data/sample_docs/",
                   "#   commercial_invoice/docs/   (100 PDFs)",
                   "#   commercial_invoice/images/ (100 PNGs @ 300 DPI)",
                   "#   bill_of_lading/docs/        (100 PDFs)",
                   "#   bill_of_lading/images/      (100 PNGs @ 300 DPI)",
                   "#   packing_list/docs/           (100 PDFs)",
                   "#   packing_list/images/         (100 PNGs @ 300 DPI)"])
    doc.add_paragraph()
    add_body(doc, "Documents use real data pools:")
    add_bullet(doc, "40 Indonesian PT company importers (Bekasi/Cikarang industrial zone)")
    add_bullet(doc, "40+ foreign exporters (Korea, China, Japan, USA, Germany)")
    add_bullet(doc, "33 real HS codes including restricted categories (weapons, chemicals)")
    add_bullet(doc, "10 shipping line container prefixes (TCKU, MSCU, MAEU, HLCU, CMAU...)")
    add_bullet(doc, "25 vessels on Asia-Indonesia trade routes")
    add_bullet(doc, "20 ports of origin worldwide")

    doc.add_paragraph()
    add_heading(doc, "6.4 Retraining with Real Data (August 2026)", level=2)
    add_body(doc, "After company visit July 27-31 to get real CEISA data:")
    add_code(doc, ["# Replace training data with real company data, then:",
                   "python scripts/train_risk_model.py",
                   "python scripts/generate_ner_training.py",
                   "python scripts/train_ner_model.py",
                   "uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload"])

    doc.add_page_break()

    # ── SECTION 7: WEB DASHBOARD ─────────────────────────────────────────────
    add_heading(doc, "7. Web Admin Dashboard")

    add_heading(doc, "7.1 Navigation Pages", level=2)
    add_table(doc,
        ["Group", "Page", "What it shows"],
        [
            ["Main", "Overview", "KPI cards (Total/Pending/Approved/Rejected/CEISA), lane distribution, recent activity"],
            ["Main", "Declarations", "Table with filters, batch select, CSV export, approve/reject workflow"],
            ["Main", "Analysis", "Risk histogram, OCR confidence, field extraction quality charts"],
            ["Main", "CEISA Portal", "All submissions with Green/Yellow/Red lane badges"],
            ["Operations", "SLA Dashboard", "Avg review time, overdue alerts (24h target), daily throughput"],
            ["Operations", "Operators", "Add/deactivate field operators, reset PINs"],
            ["Operations", "Watchlist", "Flag importers, exporters, HS codes for automatic risk elevation"],
            ["Operations", "Risk Rules", "Configure field + condition + value + boost_points rules"],
            ["Operations", "Audit Trail", "Every action: who did what, when, on which declaration"],
        ],
        col_widths=[1.0, 1.3, 3.8]
    )

    doc.add_paragraph()
    add_heading(doc, "7.2 Declaration Detail Panel", level=2)
    add_body(doc, "Clicking any row opens a right slide-in panel with:")
    features = [
        "Original document image (fetched from DB on demand)",
        "Risk score bar + percentage",
        "SHAP explanation bar chart — shows which features drove the score",
        "Extraction method badge (spaCy NER / NER+Regex / Regex)",
        "All 11 extracted fields — editable inline",
        "Flagged issues list",
        "Approve / Reject buttons with optional manager note (sends FCM push to operator)",
        "Re-run OCR button — reprocesses the stored image",
        "Submit to CEISA button (only enabled when ceisa_ready = true)",
    ]
    for f in features:
        add_bullet(doc, f)

    doc.add_paragraph()
    add_heading(doc, "7.3 Manager Login", level=2)
    add_body(doc, "URL: http://localhost:3000     Login: manager / flashport2026")

    doc.add_page_break()

    # ── SECTION 8: DATABASE ──────────────────────────────────────────────────
    add_heading(doc, "8. Database Schema")

    add_heading(doc, "8.1 Core Tables", level=2)
    add_table(doc,
        ["Table", "Key columns"],
        [
            ["declarations", "id, scan_id, document_type, operator_id, image_data (base64), tesseract_text, confidence_badge, risk_score, risk_badge, flagged_fields, ceisa_ready, review_status, review_note, reviewed_by, reviewed_at"],
            ["declaration_fields", "declaration_id, field_name, field_value, is_edited, edit_source"],
            ["ceisa_submissions", "declaration_id, jalur, response_code, response_message, submitted_by, submitted_at"],
            ["sync_logs", "scan_id, device_id, image_size_bytes, sync_duration_ms, status"],
            ["operators", "employee_id, name, pin_hash, is_active, created_at"],
        ],
        col_widths=[1.5, 4.6]
    )

    doc.add_paragraph()
    add_heading(doc, "8.2 Operations Tables (Added June 2026)", level=2)
    add_table(doc,
        ["Table", "Purpose", "Key columns"],
        [
            ["audit_logs", "Every action logged", "action, entity_type, entity_id, performed_by, detail (JSON), created_at"],
            ["watchlist", "Flagged entities", "entity_type (importer/exporter/hs_code), value, reason, is_active"],
            ["risk_rules", "Custom scoring rules", "field, condition, value, risk_boost, flag_label, is_active"],
            ["hs_code_reference", "HS validation DB", "code, description, category, is_restricted, restriction_note"],
        ],
        col_widths=[1.3, 1.4, 3.4]
    )

    doc.add_paragraph()
    add_heading(doc, "8.3 Declaration Review Status Flow", level=2)
    add_code(doc, ["pending  -->  approved  (FCM push sent to operator)",
                   "         -->  rejected  (FCM push sent with manager note)",
                   "         <--  pending   (manager can reset back)"])

    doc.add_page_break()

    # ── SECTION 9: API ───────────────────────────────────────────────────────
    add_heading(doc, "9. API Reference (36 Endpoints)")
    add_table(doc,
        ["Group", "Endpoints"],
        [
            ["Auth", "POST /auth/login,  POST /auth/operator/login"],
            ["Sync", "POST /sync  — full pipeline: detect type -> OCR -> NER -> XGBoost -> store -> broadcast"],
            ["Declarations", "GET,  GET /{id}/image,  PATCH /{id}/field,  PATCH /{id}/review,  POST /{id}/reprocess"],
            ["CEISA", "POST /ceisa/submit,  POST /ceisa/batch-submit,  GET /ceisa/submissions"],
            ["Operators", "GET, POST, PATCH /{id}, POST /{id}/reset-pin, DELETE /{id}"],
            ["Watchlist", "GET, POST, DELETE /{id}, GET /check"],
            ["Risk Rules", "GET, POST, PATCH /{id}, DELETE /{id}"],
            ["HS Codes", "GET, POST, GET /validate/{code}, DELETE /{code}"],
            ["SLA", "GET /sla  — avg review time, overdue alerts, daily throughput"],
            ["Audit", "GET /audit  — filter by entity_id, entity_type"],
            ["Export", "GET /export/declarations.csv  — with status and date filters"],
            ["Health", "GET /health"],
        ],
        col_widths=[1.4, 4.7]
    )

    doc.add_page_break()

    # ── SECTION 10: RUNNING ──────────────────────────────────────────────────
    add_heading(doc, "10. Running the System")

    add_heading(doc, "10.1 Prerequisites", level=2)
    add_bullet(doc, "PostgreSQL 15 (local install via Homebrew — no Docker required)")
    add_bullet(doc, "Python 3.11 + pip")
    add_bullet(doc, "Node.js 18+")
    add_bullet(doc, "Flutter 3.x")
    add_bullet(doc, "Tesseract OCR: brew install tesseract")

    doc.add_paragraph()
    add_heading(doc, "10.2 Start Commands", level=2)
    add_body(doc, "Backend:", bold=True)
    add_code(doc, ["cd backend && source venv/bin/activate",
                   "uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload"])
    doc.add_paragraph()
    add_body(doc, "Web Dashboard:", bold=True)
    add_code(doc, ["cd web && npm run dev",
                   "# Opens at http://localhost:3000",
                   "# Login: manager / flashport2026"])
    doc.add_paragraph()
    add_body(doc, "Mobile (simulator):", bold=True)
    add_code(doc, ["cd mobile && flutter run"])
    doc.add_paragraph()
    add_body(doc, "Mobile (physical device):", bold=True)
    add_code(doc, ["# 1. Find Mac IP:  ipconfig getifaddr en0",
                   "# 2. In app login screen -> Server URL -> http://[IP]:8000",
                   "# 3. Login: CDP-001 / 1234"])

    doc.add_paragraph()
    add_heading(doc, "10.3 Seeded Operator Accounts", level=2)
    add_table(doc,
        ["Employee ID", "Name", "PIN"],
        [
            ["CDP-001", "Ahmad Fauzi", "1234"],
            ["CDP-002", "Budi Santoso", "5678"],
            ["CDP-003", "Citra Dewi", "9012"],
        ],
        col_widths=[1.5, 2.0, 1.0]
    )

    doc.add_page_break()

    # ── SECTION 11: PHASE 2 ──────────────────────────────────────────────────
    add_heading(doc, "11. Phase 2 — August 2026")
    add_table(doc,
        ["Item", "Action"],
        [
            ["Real CEISA H2H integration",
             "After company visit July 27-31. Replace mock gateway in app/core/ceisa_gateway.py with real HTTP calls. Add CEISA_API_URL, CEISA_CLIENT_ID, CEISA_SECRET to .env"],
            ["XGBoost retrain",
             "Feed real CEISA rejection data into data/training_declarations.csv then run python scripts/train_risk_model.py"],
            ["NER retrain",
             "Add real document scans to sample_docs/ then run generate_ner_training.py + train_ner_model.py"],
            ["FCM activation",
             "Firebase console -> Project Settings -> Service Accounts -> Generate new private key. Set FCM_SERVICE_ACCOUNT_JSON in backend/.env"],
        ],
        col_widths=[2.0, 4.1]
    )

    doc.add_paragraph()

    # ── SECTION 12: FILE STRUCTURE ───────────────────────────────────────────
    add_heading(doc, "12. Project File Structure")
    add_code(doc, [
        "flashport/",
        "  docs/                          All documentation",
        "    FlashPort_Proposal.docx",
        "    FlashPort_System_Breakdown_Updated.docx  (this file)",
        "    FlashPort_How_It_Works.pdf",
        "    Logistic - Cikarang Dryport.pdf",
        "",
        "  backend/",
        "    app/",
        "      api/         14 route files, 36 total endpoints",
        "      models/      7 SQLAlchemy models",
        "      services/    extractor, risk_scorer, ocr, preprocessing, fcm, audit",
        "      core/        ceisa_gateway (mock Phase 1)",
        "    scripts/       6 training + generation scripts",
        "    models/        risk_model.xgb, ner_model/, model_info.json",
        "    data/",
        "      training_declarations.csv  (600 records)",
        "      ner_training.spacy         (300 annotated examples)",
        "      sample_docs/               (300 PDFs + 300 PNGs)",
        "",
        "  web/src/",
        "    components/    15 React components",
        "    hooks/         4 custom hooks",
        "",
        "  mobile/lib/",
        "    screens/       5 screens (home, camera, preview, result, login)",
        "    services/      4 services (sync, database, operator, backend_config)",
        "    widgets/       scan_tile with status timeline",
        "",
        "  docker/postgres/init.sql       Full DB schema",
        "  README.md",
        "  CLAUDE.md                      AI developer context",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "FlashPort  |  AI Open Innovation Challenge 2026  |  "
        "Team Teknik Logistik  |  President University  |  Updated June 2026"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved -> {OUT}")
    print(f"Size:  {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
